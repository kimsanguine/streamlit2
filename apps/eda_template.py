# 미니프로젝트 연동 — 범용 CSV EDA 대시보드
# [왜] AI Python 2 미니프로젝트(heart.csv 등)는 노트북 안에서만 끝난다. 같은 CSV를 이 템플릿에
#      올리면 "내 데이터가 웹 대시보드가 되는" 경험을 3분 안에 체감할 수 있다 — ST2 penguins
#      대시보드(apps/m3_penguins.py)와 같은 탭 구조를 범용 CSV로 확장한 버전이다.
# [흐름] 미니프로젝트_5_Kaggle_Heart_Disease_분류(정답_light).ipynb의 Word 보고서(build_report) 관례를
#      계승 — EDA 결과를 docx로 내려받는 기능을 이 템플릿에도 넣는다.
# 실행: python3.11 -m streamlit run apps/eda_template.py

import io

# [흐름] 4개 라이브러리로 충분 — 데이터(pandas)·데모 데이터셋(seaborn)·시각화(matplotlib)·웹 UI(streamlit)
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import streamlit as st

# [왜] matplotlib 기본 폰트(DejaVu Sans)엔 한글 글리프가 없어 라벨이 □□로 깨진다.
#      로컬은 Mac(AppleGothic)이 있어 멀쩡하지만, Streamlit Cloud(Linux)엔 한글 폰트가 없어 배포 때 깨진다.
#      → packages.txt의 fonts-nanum으로 Cloud에 NanumGothic을 설치하고, 아래에서 설치된 폰트를 순서대로 찾는다.
from matplotlib import font_manager

for _f in ["AppleGothic", "Malgun Gothic", "NanumGothic", "NanumBarunGothic"]:
    if any(_font.name == _f for _font in font_manager.fontManager.ttflist):
        plt.rcParams["font.family"] = _f
        break
# [왜] 마이너스 기호(−)는 한글 폰트에 글리프가 없는 경우가 많다 — unicode_minus=False로 일반 하이픈을 쓴다
plt.rcParams["axes.unicode_minus"] = False

# [교안 팔레트] 강사님 교안과 통일 — 이산 색이 필요한 자리(히스토그램·타깃별 박스플롯)에 쓴다.
# [흐름] NAVY는 분포 탭의 히스토그램, TEAL·CORAL은 타깃 탭의 박스플롯에 각각 쓰인다
NAVY = "#1F4E79"
TEAL = "#0E6B56"
CORAL = "#993C1D"


# [왜] cache_data — 업로드 파일은 rerun마다 다시 파싱하면 느리다. file_bytes(불변 바이트)를 키로
#      캐싱해, 같은 파일이면 재파싱 없이 즉시 재사용한다(ST1에서 배운 그 데코레이터).
@st.cache_data
def load_data(file_bytes: bytes | None, filename: str | None) -> tuple[pd.DataFrame, str]:
    # [흐름] 아직 아무 파일도 안 올렸으면(file_bytes=None) 데모 데이터(팔머펭귄)로 화면을 채운다 —
    #      "빈 화면"보다 "동작하는 예시"가 먼저 보이는 편이 처음 켰을 때 더 안심된다.
    if file_bytes is None:
        return sns.load_dataset("penguins").dropna(), "penguins (데모 데이터)"
    # [흐름] 업로드된 CSV는 메모리 안의 바이트(BytesIO)를 그대로 읽는다 — 디스크에 임시 저장할 필요가 없다.
    # [왜] 한국어 Excel이 저장한 CSV는 cp949 인코딩이 흔하다 — UTF-8(utf-8-sig는 BOM까지 처리)로
    #      읽다 실패하면 cp949로 한 번 더 시도해, 흔한 업로드가 인코딩 오류로 죽지 않게 한다.
    try:
        return pd.read_csv(io.BytesIO(file_bytes), encoding="utf-8-sig"), filename
    except UnicodeDecodeError:
        return pd.read_csv(io.BytesIO(file_bytes), encoding="cp949"), filename


def _top_corr_pairs(numeric_df: pd.DataFrame, n: int = 3) -> list[tuple[str, str, float]]:
    """[통계] 수치형 컬럼 쌍 중 절대상관 상위 n개. st 의존 없는 순수 함수 — 테스트·docx 양쪽에서 재사용."""
    # [왜] 방향(+/-)보다 "얼마나 강한지"가 궁금하므로 절댓값(abs)으로 본다
    corr = numeric_df.corr().abs()
    cols = corr.columns
    # [흐름] i<j 조합만 순회 — (A,B)와 (B,A)는 같은 상관값이라 중복 계산을 피한다
    pairs = [
        (cols[i], cols[j], float(corr.iloc[i, j]))
        for i in range(len(cols))
        for j in range(i + 1, len(cols))
    ]
    # [흐름] 상관 강도 내림차순 정렬 — 가장 눈여겨볼 쌍이 리스트 맨 앞에 오게 한다
    pairs.sort(key=lambda x: x[2], reverse=True)
    return pairs[:n]


def _build_report_docx(df: pd.DataFrame, source_name: str, target_col: str | None) -> bytes:
    """[흐름] AI Python 2 미니프로젝트의 build_report() 관례 계승 — 화면에서 이미 계산한 값을
    그대로 옮겨 적을 뿐 재계산은 하지 않는다. python-docx 미설치 시 호출되지 않도록 main()에서 방어."""
    # [왜] 여기서만 import — python-docx가 없어도 이 함수를 안 부르면 앱 자체는 안 죽는다
    from docx import Document

    # [흐름] 문서 맨 위 — 출처·행수·열수로 이 보고서가 어떤 데이터를 다뤘는지 먼저 밝힌다
    doc = Document()
    doc.add_heading("EDA 분석보고서", level=1)
    doc.add_paragraph(f"데이터 출처: {source_name}")
    doc.add_paragraph(f"행 수: {len(df)}  ·  열 수: {len(df.columns)}")

    # [흐름] 결측치 섹션 — 있으면 표로, 없으면 문장으로. 화면(tab_target)과 같은 계산을 문서에도 반영한다.
    doc.add_heading("결측치 요약", level=2)
    missing = df.isna().sum()
    missing = missing[missing > 0].sort_values(ascending=False)
    if len(missing):
        # [흐름] docx 표는 Streamlit dataframe과 달리 행을 하나씩 add_row()로 직접 채워야 한다
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "컬럼"
        table.rows[0].cells[1].text = "결측 개수"
        for col, cnt in missing.items():
            row = table.add_row().cells
            row[0].text = str(col)
            row[1].text = str(int(cnt))
    else:
        doc.add_paragraph("결측치 없음")

    # [흐름] 상관관계 Top 3 — _top_corr_pairs()로 계산한 순수 함수 결과를 그대로 문장으로 옮긴다
    doc.add_heading("상관관계 Top 3", level=2)
    numeric_df = df.select_dtypes("number")
    if numeric_df.shape[1] >= 2:
        for a, b, r in _top_corr_pairs(numeric_df, n=3):
            doc.add_paragraph(f"{a} ↔ {b} : r = {r:.3f}", style="List Bullet")
    else:
        doc.add_paragraph("수치형 컬럼이 2개 미만이라 상관관계를 계산할 수 없음")

    # [흐름] 타깃 컬럼을 골랐을 때만 값 분포까지 문서에 남긴다 — 안 고르면 이 섹션 자체가 생략된다
    if target_col:
        doc.add_heading("타깃 컬럼", level=2)
        # [흐름] value_counts()를 dict()로 감싸 "값: 개수" 형태 그대로 문장에 옮긴다
        doc.add_paragraph(f"{target_col} — 값 분포: {dict(df[target_col].value_counts())}")

    # [흐름] 파일로 저장하지 않고 메모리(BytesIO)에 담아 바이트로 반환 — st.download_button이 바로 쓴다
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    # [왜] set_page_config는 다른 st.* 호출보다 먼저 와야 한다 — 탭 제목·아이콘·레이아웃을 정하는 설정
    st.set_page_config(page_title="CSV EDA 대시보드", page_icon="📊", layout="wide")
    st.title("📊 CSV EDA 대시보드")
    # [왜] caption으로 이 앱의 쓰임새를 3초 안에 요약 — 제목만으론 "무엇을 3분 안에 하는지" 안 드러난다
    st.caption("내 CSV를 올려 미리보기·분포·상관·타깃 분포까지 3분 안에 훑어보는 범용 템플릿")

    # [원칙③] 선수 지식 연결 박스 — 이 템플릿이 ST2 penguins·AI Python 2와 어떻게 이어지는지 항상 먼저 보여준다
    with st.expander("💡 이전 챕터와 연결"):
        st.markdown(
            "- **ST2 penguins 대시보드**(`apps/m3_penguins.py`)와 같은 탭 구조(미리보기 → 시각화 → "
            "심화)를 임의의 CSV로 확장한 버전입니다.\n"
            "- **AI Python 2 미니프로젝트**에서 `df.describe()`·`df.isna().sum()`·상관 히트맵으로 "
            "손으로 하던 EDA가, 여기서는 파일 업로드 한 번으로 재현됩니다."
        )

    # [흐름] file_uploader는 파일이 없으면 None을 반환한다 — getvalue()/name도 그 None 분기를 따라간다
    # [왜] type=["csv"] — 다른 확장자를 올리면 업로드 창 단계에서 미리 걸러 pandas 에러를 막는다
    uploaded = st.file_uploader(
        "CSV 업로드 (없으면 데모 데이터로 동작 — 팔머펭귄 데이터셋)", type=["csv"]
    )
    # [왜] uploaded 객체 자체가 아니라 getvalue()(불변 바이트)를 캐시 키로 넘긴다 — cache_data는
    #      인자가 "같은 값"이어야 캐시를 재사용하는데, UploadedFile 객체는 매 rerun마다 새로 만들어진다.
    file_bytes = uploaded.getvalue() if uploaded is not None else None
    filename = uploaded.name if uploaded is not None else None
    df, source_name = load_data(file_bytes, filename)
    # [왜] 지금 화면이 데모 데이터인지 내 CSV인지 항상 표시 — 헷갈리면 잘못된 데이터로 분석을 이어가기 쉽다
    st.caption(f"현재 데이터: **{source_name}**")

    # [왜] 탭 4개 — "다른 질문"을 화면 하나에 다 욱여넣지 않고 목적별로 분리한다(m3_penguins.py와 같은 원칙).
    tab_preview, tab_dist, tab_corr, tab_target = st.tabs(
        ["🔍 미리보기", "📈 분포", "🔥 상관 히트맵", "🎯 결측·타깃"]
    )

    with tab_preview:
        # [흐름] Shape·dtypes·head 3가지 — "이 데이터가 뭔지" 감을 잡는 가장 기본적인 첫 확인 3종
        st.subheader("Shape")
        # [왜] columns(2) — 행·열 숫자를 나란히 둬 Shape을 한눈에 파악하게 한다
        c1, c2 = st.columns(2)
        # [왜] metric — 숫자 하나를 표보다 크고 눈에 띄게 보여줄 때 쓴다(m3_penguins.py와 같은 위젯)
        c1.metric("행 (rows)", df.shape[0])
        c2.metric("열 (columns)", df.shape[1])

        st.subheader("컬럼 타입 (dtypes)")
        # [왜] astype(str) — dtype 객체를 그대로 넘기면 표에 "dtype('int64')"처럼 안 예쁘게 나온다
        st.dataframe(df.dtypes.astype(str).rename("dtype"))

        st.subheader("데이터 미리보기 (head)")
        # [왜] head() — 전체를 다 그리면 큰 CSV에서 화면이 느려진다, 상위 5행만으로도 컬럼 감각은 충분하다
        st.dataframe(df.head())

    # [흐름] tab_dist·tab_corr·tab_target 3곳에서 공통으로 쓰는 수치형 컬럼 목록 — 한 번만 계산해 재사용한다
    numeric_cols = df.select_dtypes("number").columns.tolist()

    with tab_dist:
        # [왜] 컬럼마다 스케일이 다르다 — 한 화면에 다 그리면 서로 뭉개진다. selectbox로 하나씩
        #      골라 보게 하고, dropna()로 결측치가 히스토그램 계산에 끼어들지 않게 막는다.
        if not numeric_cols:
            st.info("수치형 컬럼이 없습니다.")
        else:
            col = st.selectbox("분포를 볼 컬럼", numeric_cols)
            fig, ax = plt.subplots()
            ax.hist(df[col].dropna(), bins=30, color=NAVY)
            # [왜] 축 이름을 명시하지 않으면 기본값(빈 문자열)이라 어떤 컬럼의 분포인지 그래프만 봐서는 알 수 없다
            ax.set_xlabel(col)
            ax.set_ylabel("빈도")
            st.pyplot(fig)
            # [왜] plt.close — 안 닫으면 rerun마다 figure가 메모리에 계속 쌓인다
            plt.close(fig)

    with tab_corr:
        # [왜] 수치형 컬럼이 1개 이하면 상관계수 자체가 정의되지 않는다 — 계산 전에 먼저 막는다.
        if len(numeric_cols) < 2:
            st.info("상관관계를 그리려면 수치형 컬럼이 2개 이상 필요합니다.")
        else:
            numeric_df = df[numeric_cols]
            # [흐름] 컬럼 수에 비례해 figure 크기를 늘리되 10인치를 상한으로 — 컬럼이 많아도 라벨이 안 겹치게
            fig, ax = plt.subplots(figsize=(min(1 + 0.6 * len(numeric_cols), 10),) * 2)
            sns.heatmap(numeric_df.corr(), annot=True, fmt=".2f", cmap="Blues", ax=ax)
            st.pyplot(fig)
            plt.close(fig)
            # [원칙④] 히트맵을 보여주는 것으로 끝내지 않고, 실무에서는 다음에 뭘 하는지 한 줄로 닫는다
            st.caption(
                "→ 실무에서는 이 히트맵으로 먼저 후보를 좁히고, 진짜 인과·유의성 판단은 "
                "미니프로젝트에서 배운 Welch t-test·Cohen's d 같은 통계 검정으로 확인합니다."
            )

    with tab_target:
        # [왜] 결측치와 타깃 분포를 한 탭에 묶었다 — 둘 다 "이 데이터를 모델에 넣기 전에
        #      확인해야 할 것"이라는 같은 목적을 가진 점검 항목이라서다.
        st.subheader("결측치")
        # [흐름] docx의 결측치 계산과 동일한 로직 — 화면·보고서가 같은 숫자를 보여줘야 신뢰가 생긴다
        missing = df.isna().sum()
        missing = missing[missing > 0].sort_values(ascending=False)
        if len(missing):
            st.dataframe(missing.rename("결측 개수"))
        else:
            st.success("결측치 없음")

        st.subheader("타깃별 분포")
        # [흐름] 타깃 컬럼은 "선택 안 함"을 기본값으로 둔다 — CSV마다 타깃 이름이 다르므로 자동 추측 대신 직접 고르게 한다
        target_col = st.selectbox("타깃 컬럼 선택 (분류 문제의 정답 열)", ["(선택 안 함)"] + list(df.columns))
        # [흐름] 타깃 컬럼을 고른 뒤에만 비교할 수치형 컬럼을 추가로 고르게 한다 — 순서가 곧 분석 흐름이다
        if target_col != "(선택 안 함)" and numeric_cols:
            feature_col = st.selectbox("비교할 수치형 컬럼", numeric_cols, key="target_feature")
            fig, ax = plt.subplots()
            # [왜] 박스플롯 — 타깃 값별로 수치형 컬럼의 분포(중앙값·사분위)가 갈리는지 한눈에 비교한다
            sns.boxplot(data=df, x=target_col, y=feature_col, hue=target_col,
                        palette=[TEAL, CORAL, NAVY], legend=False, ax=ax)
            st.pyplot(fig)
            plt.close(fig)
        elif target_col != "(선택 안 함)":
            # [흐름] 타깃은 골랐지만 수치형 컬럼이 하나도 없는 경우(전부 범주형) — 박스플롯을 못 그린다는 안내만 준다
            st.info("비교할 수치형 컬럼이 없습니다.")

    # [흐름] divider로 지금까지의 탐색 파트와 보고서 다운로드 파트를 시각적으로 분리한다
    st.divider()
    st.subheader("📄 분석보고서 다운로드")
    # [왜] python-docx는 requirements.txt에 있지만, 별도 환경에서 누락됐을 때 앱 전체가
    #      죽지 않도록 try/except로 방어하고, 없으면 설치 안내만 보여준다.
    try:
        # noqa: F401  — 설치 여부만 확인
        from docx import Document
        docx_available = True
    except ImportError:
        docx_available = False

    if docx_available:
        # [흐름] tab_target에서 고른 target_col을 그대로 재사용 — 같은 선택을 두 번 하게 만들지 않는다
        target_for_report = target_col if target_col != "(선택 안 함)" else None
        report_bytes = _build_report_docx(df, source_name, target_for_report)
        # [왜] mime 타입을 정확히 지정해야 브라우저가 확장자 없는 파일로 받지 않고 .docx로 인식한다
        st.download_button(
            "📄 EDA 요약 docx 다운로드",
            data=report_bytes,
            file_name="EDA_분석보고서.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        # [왜] 버튼을 아예 숨기지 않고 설치 방법을 안내 — 사용자가 "왜 기능이 없는지" 알 수 있게 한다
        st.info("docx 보고서를 쓰려면 터미널에서 `pip install python-docx`를 설치하세요.")


# [왜] if __name__ == "__main__": 가드 — streamlit run이 이 파일을 직접 실행할 때만 UI가 뜨고,
#      나중에 다른 스크립트가 이 파일을 import해도 화면이 저절로 뜨지 않도록 막아준다.
if __name__ == "__main__":
    main()
